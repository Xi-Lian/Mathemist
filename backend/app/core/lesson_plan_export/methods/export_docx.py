from .._shared import *


class _ExportDocxMixin:
    def export_docx(
        self,
        lesson_plan_content: str,
        filename: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> str:
        """
        导出为Word文档（DOCX）格式
        
        Args:
            lesson_plan_content: 教案内容（Markdown格式）
            filename: 文件名（不含扩展名）
            metadata: 元数据字典
        
        Returns:
            导出的文件路径
        """
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("请安装python-docx库: pip install python-docx")
        
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"lesson_plan_{timestamp}"
        
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        
        if metadata and 'topic' in metadata:
            title = doc.add_heading(metadata['topic'], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            title = doc.add_heading('教案', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if metadata:
            meta_para = doc.add_paragraph()
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            meta_text = []
            if 'chapter' in metadata:
                meta_text.append(f"章节：{metadata['chapter']}")
            if 'textbook' in metadata:
                meta_text.append(f"教材：{metadata['textbook']}")
            if meta_text:
                meta_para.add_run(' | '.join(meta_text))
                meta_para.add_run().add_break()
        
        doc.add_paragraph()
        self._add_markdown_to_docx(doc, lesson_plan_content)
        filepath = self.output_dir / f"{filename}.docx"
        doc.save(str(filepath))
        
        print(f"✅ Word文档导出成功: {filepath}")
        return str(filepath)
