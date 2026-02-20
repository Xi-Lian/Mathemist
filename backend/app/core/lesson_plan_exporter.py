"""
教案导出模块

职责：
- 将生成的教案导出为多种格式（Markdown、HTML、DOCX、PDF）
- 支持自定义导出模板
- 保持教案的格式和结构

依赖：
- markdown (Markdown处理)
- python-docx (Word文档生成)
- weasyprint (PDF生成，可选)
- pathlib (路径管理)
"""

import os
from pathlib import Path
from typing import Dict, Any, Optional
from datetime import datetime
import markdown


class LessonPlanExporter:
    """教案导出器"""
    
    def __init__(self, output_dir: Optional[str] = None):
        """
        初始化教案导出器
        
        Args:
            output_dir: 输出目录，默认为 backend/exports
        """
        if output_dir:
            self.output_dir = Path(output_dir)
        else:
            self.output_dir = Path(__file__).parent.parent.parent / "backend" / "exports"
        
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def export_markdown(
        self,
        lesson_plan_content: str,
        filename: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> str:
        """
        导出为Markdown格式
        
        Args:
            lesson_plan_content: 教案内容（Markdown格式）
            filename: 文件名（不含扩展名）
            metadata: 元数据字典
        
        Returns:
            导出的文件路径
        """
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"lesson_plan_{timestamp}"
        
        content = self._add_metadata_header(lesson_plan_content, metadata)
        filepath = self.output_dir / f"{filename}.md"
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        
        print(f"✅ Markdown导出成功: {filepath}")
        return str(filepath)
    
    def export_html(
        self,
        lesson_plan_content: str,
        filename: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None,
        include_css: bool = True
    ) -> str:
        """
        导出为HTML格式
        
        Args:
            lesson_plan_content: 教案内容（Markdown格式）
            filename: 文件名（不含扩展名）
            metadata: 元数据字典
            include_css: 是否包含CSS样式
        
        Returns:
            导出的文件路径
        """
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"lesson_plan_{timestamp}"
        
        html_content = markdown.markdown(
            lesson_plan_content,
            extensions=['extra', 'tables', 'toc']
        )
        html_doc = self._build_html_document(html_content, metadata, include_css)
        filepath = self.output_dir / f"{filename}.html"
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html_doc)
        
        print(f"✅ HTML导出成功: {filepath}")
        return str(filepath)
    
    def export_docx(
        self,
        lesson_plan_content: str,
        filename: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> str:
        """
        导出为Word文档（DOCX）格式
        
        Args:
            lesson_plan_content: 教案内容（Markdown格式）
            filename: 文件名（不含扩展名）
            metadata: 元数据字典
        
        Returns:
            导出的文件路径
        """
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("请安装python-docx库: pip install python-docx")
        
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"lesson_plan_{timestamp}"
        
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        
        if metadata and 'topic' in metadata:
            title = doc.add_heading(metadata['topic'], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            title = doc.add_heading('教案', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if metadata:
            meta_para = doc.add_paragraph()
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            meta_text = []
            if 'chapter' in metadata:
                meta_text.append(f"章节：{metadata['chapter']}")
            if 'textbook' in metadata:
                meta_text.append(f"教材：{metadata['textbook']}")
            if meta_text:
                meta_para.add_run(' | '.join(meta_text))
                meta_para.add_run().add_break()
        
        doc.add_paragraph()
        self._add_markdown_to_docx(doc, lesson_plan_content)
        filepath = self.output_dir / f"{filename}.docx"
        doc.save(str(filepath))
        
        print(f"✅ Word文档导出成功: {filepath}")
        return str(filepath)
    
    def export_pdf(
        self,
        lesson_plan_content: str,
        filename: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> str:
        """
        导出为PDF格式
        
        Args:
            lesson_plan_content: 教案内容（Markdown格式）
            filename: 文件名（不含扩展名）
            metadata: 元数据字典
        
        Returns:
            导出的文件路径
        """
        try:
            from weasyprint import HTML
        except ImportError:
            raise ImportError("请安装weasyprint库: pip install weasyprint")
        
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"lesson_plan_{timestamp}"
        
        html_content = markdown.markdown(
            lesson_plan_content,
            extensions=['extra', 'tables', 'toc']
        )
        html_doc = self._build_html_document(html_content, metadata, include_css=True)
        filepath = self.output_dir / f"{filename}.pdf"
        HTML(string=html_doc).write_pdf(str(filepath))
        
        print(f"✅ PDF导出成功: {filepath}")
        return str(filepath)
    
    def export_all(
        self,
        lesson_plan_content: str,
        filename: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None,
        formats: Optional[list] = None
    ) -> Dict[str, str]:
        """
        导出为多种格式
        
        Args:
            lesson_plan_content: 教案内容（Markdown格式）
            filename: 文件名（不含扩展名）
            metadata: 元数据字典
            formats: 要导出的格式列表，默认为 ['markdown', 'html', 'docx']
        
        Returns:
            格式到文件路径的字典
        """
        if formats is None:
            formats = ['markdown', 'html', 'docx']
        
        results = {}
        for format_type in formats:
            try:
                if format_type == 'markdown':
                    results['markdown'] = self.export_markdown(
                        lesson_plan_content, filename, metadata
                    )
                elif format_type == 'html':
                    results['html'] = self.export_html(
                        lesson_plan_content, filename, metadata
                    )
                elif format_type == 'docx':
                    results['docx'] = self.export_docx(
                        lesson_plan_content, filename, metadata
                    )
                elif format_type == 'pdf':
                    results['pdf'] = self.export_pdf(
                        lesson_plan_content, filename, metadata
                    )
            except Exception as e:
                print(f"⚠️  导出{format_type}格式失败: {e}")
                results[format_type] = f"导出失败: {e}"
        
        return results
    
    def _add_metadata_header(
        self,
        content: str,
        metadata: Optional[Dict[str, Any]]
    ) -> str:
        """
        添加元数据头部
        
        Args:
            content: 原始内容
            metadata: 元数据字典
        
        Returns:
            添加元数据后的内容
        """
        if not metadata:
            return content
        
        header_lines = ["---"]
        for key, value in metadata.items():
            header_lines.append(f"{key}: {value}")
        header_lines.append(f"export_time: {datetime.now().isoformat()}")
        header_lines.append("---")
        header_lines.append("")
        
        return "\n".join(header_lines) + content
    
    def _build_html_document(
        self,
        html_content: str,
        metadata: Optional[Dict[str, Any]],
        include_css: bool
    ) -> str:
        """
        构建完整的HTML文档
        
        Args:
            html_content: HTML内容
            metadata: 元数据字典
            include_css: 是否包含CSS
        
        Returns:
            完整的HTML文档
        """
        title = "教案"
        if metadata and 'topic' in metadata:
            title = metadata['topic']
        
        css = self._get_css() if include_css else ""
        meta_tags = ""
        if metadata:
            for key, value in metadata.items():
                meta_tags += f'<meta name="{key}" content="{value}">\n'
        
        html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    {meta_tags}
    <title>{title}</title>
    {css}
</head>
<body>
    <div class="container">
        {html_content}
    </div>
</body>
</html>"""
        return html
    
    def _get_css(self) -> str:
        """获取CSS样式"""
        return """<style>
    * {
        margin: 0;
        padding: 0;
        box-sizing: border-box;
    }
    body {
        font-family: "Microsoft YaHei", "SimSun", sans-serif;
        line-height: 1.8;
        color: #333;
        background-color: #f5f5f5;
        padding: 20px;
    }
    .container {
        max-width: 900px;
        margin: 0 auto;
        background-color: white;
        padding: 40px;
        box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        border-radius: 8px;
    }
    h1 {
        text-align: center;
        color: #2c3e50;
        margin-bottom: 30px;
        padding-bottom: 15px;
        border-bottom: 2px solid #3498db;
    }
    h2 {
        color: #34495e;
        margin-top: 30px;
        margin-bottom: 15px;
        padding-left: 10px;
        border-left: 4px solid #3498db;
    }
    h3 {
        color: #555;
        margin-top: 20px;
        margin-bottom: 10px;
    }
    p {
        margin-bottom: 15px;
        text-align: justify;
    }
    ul, ol {
        margin-left: 30px;
        margin-bottom: 15px;
    }
    li {
        margin-bottom: 8px;
    }
    table {
        width: 100%;
        border-collapse: collapse;
        margin: 20px 0;
    }
    th, td {
        border: 1px solid #ddd;
        padding: 12px;
        text-align: left;
    }
    th {
        background-color: #3498db;
        color: white;
        font-weight: bold;
    }
    tr:nth-child(even) {
        background-color: #f9f9f9;
    }
    tr:hover {
        background-color: #f5f5f5;
    }
    code {
        background-color: #f4f4f4;
        padding: 2px 6px;
        border-radius: 3px;
        font-family: "Courier New", monospace;
    }
    pre {
        background-color: #f4f4f4;
        padding: 15px;
        border-radius: 5px;
        overflow-x: auto;
        margin: 15px 0;
    }
    blockquote {
        border-left: 4px solid #3498db;
        padding-left: 20px;
        margin: 20px 0;
        color: #666;
        font-style: italic;
    }
    @media print {
        body {
            background-color: white;
            padding: 0;
        }
        .container {
            box-shadow: none;
            padding: 20px;
        }
    }
</style>"""
    
    def _add_markdown_to_docx(self, doc, markdown_content: str):
        """
        将Markdown内容添加到Word文档
        
        Args:
            doc: Word文档对象
            markdown_content: Markdown内容
        """
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        lines = markdown_content.split('\n')
        in_code_block = False
        code_block_lines = []
        
        for line in lines:
            line = line.rstrip()
            if line.startswith('```'):
                if in_code_block:
                    if code_block_lines:
                        para = doc.add_paragraph('\n'.join(code_block_lines))
                        para.style = 'Intense Quote'
                    code_block_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                continue
            
            if in_code_block:
                code_block_lines.append(line)
                continue
            
            if line.startswith('# '):
                p = doc.add_heading(line[2:], level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            elif not line:
                doc.add_paragraph()
            elif line.startswith('> '):
                p = doc.add_paragraph(line[2:])
                p.style = 'Quote'
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line and line[0].isdigit() and '. ' in line:
                doc.add_paragraph(line.split('. ', 1)[1], style='List Number')
            elif line:
                doc.add_paragraph(line)


lesson_plan_exporter = LessonPlanExporter()


def export_lesson_plan_markdown(
    lesson_plan_content: str,
    filename: Optional[str] = None,
    metadata: Optional[Dict[str, Any]] = None
) -> str:
    """导出为Markdown"""
    return lesson_plan_exporter.export_markdown(lesson_plan_content, filename, metadata)


def export_lesson_plan_html(
    lesson_plan_content: str,
    filename: Optional[str] = None,
    metadata: Optional[Dict[str, Any]] = None
) -> str:
    """导出为HTML"""
    return lesson_plan_exporter.export_html(lesson_plan_content, filename, metadata)


def export_lesson_plan_docx(
    lesson_plan_content: str,
    filename: Optional[str] = None,
    metadata: Optional[Dict[str, Any]] = None
) -> str:
    """导出为Word文档"""
    return lesson_plan_exporter.export_docx(lesson_plan_content, filename, metadata)


def export_lesson_plan_pdf(
    lesson_plan_content: str,
    filename: Optional[str] = None,
    metadata: Optional[Dict[str, Any]] = None
) -> str:
    """导出为PDF"""
    return lesson_plan_exporter.export_pdf(lesson_plan_content, filename, metadata)


def export_lesson_plan_all(
    lesson_plan_content: str,
    filename: Optional[str] = None,
    metadata: Optional[Dict[str, Any]] = None,
    formats: Optional[list] = None
) -> Dict[str, str]:
    """导出为多种格式"""
    return lesson_plan_exporter.export_all(lesson_plan_content, filename, metadata, formats)
