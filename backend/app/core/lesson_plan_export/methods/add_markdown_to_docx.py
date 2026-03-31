from .._shared import *


class _AddMarkdownToDocxMixin:
    def _add_markdown_to_docx(self, doc, markdown_content: str):
        """
        将Markdown内容添加到Word文档
        
        Args:
            doc: Word文档对象
            markdown_content: Markdown内容
        """
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        lines = markdown_content.split('\n')
        in_code_block = False
        code_block_lines = []
        
        for line in lines:
            line = line.rstrip()
            if line.startswith('```'):
                if in_code_block:
                    if code_block_lines:
                        para = doc.add_paragraph('\n'.join(code_block_lines))
                        para.style = 'Intense Quote'
                    code_block_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                continue
            
            if in_code_block:
                code_block_lines.append(line)
                continue
            
            if line.startswith('# '):
                p = doc.add_heading(line[2:], level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            elif not line:
                doc.add_paragraph()
            elif line.startswith('> '):
                p = doc.add_paragraph(line[2:])
                p.style = 'Quote'
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line and line[0].isdigit() and '. ' in line:
                doc.add_paragraph(line.split('. ', 1)[1], style='List Number')
            elif line:
                doc.add_paragraph(line)
